"""Build MART_DEV_OTA_ANALYST_GUIDE.docx from structured content.

Usage:
  python data-eng/scripts/build_mart_ota_analyst_guide_docx.py
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

ROOT = Path(__file__).resolve().parents[1]
DEFAULT_OUT = ROOT / "docs" / "MART_DEV_OTA_ANALYST_GUIDE.docx"


def _shade_header_row(table) -> None:
    for cell in table.rows[0].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = tc_pr.makeelement(qn("w:shd"), {})
            tc_pr.append(shd)
        shd.set(qn("w:fill"), "E8E8E8")


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _add_body(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(11)


def _add_code(doc: Document, text: str) -> None:
    for line in text.strip("\n").splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(0)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], value)
    _shade_header_row(table)
    doc.add_paragraph()


def _heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("Mart Dev Database Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    sub = doc.add_paragraph("For OTA Insights Analysts")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in sub.runs:
        run.italic = True
        run.font.size = Pt(14)

    for line in [
        "Prepared by: OpenTrace Data Team",
        "Date: August 2026",
        "Dataset: mart_dev (BigQuery analytics-ready)",
        "Companion: mart_dev_entity_dictionary.xlsx",
        "Classification: Internal — for analysts authoring OTA insights",
    ]:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    _heading(doc, "1. Purpose of this guide")
    _add_body(
        doc,
        "This document teaches analysts how to use the OpenTrace mart_dev warehouse when writing "
        "OTA insights analytical reports (insight / metric / recommendation narratives).",
    )
    _add_body(
        doc,
        "OTA insights are authored analytical products. Metrics and claims should be grounded in "
        "mart_dev (and may later be ingested into the Ask ADZA OTA_insights corpus). Use the Excel "
        "entity dictionary for per-table detail; use this guide for how to think and query.",
    )

    _heading(doc, "2. How OTA insights relate to mart_dev")
    _add_table(
        doc,
        ["Layer", "Role"],
        [
            ["mart_dev", "Star-schema facts, dimensions, aggregates — source of truth for numbers"],
            ["OTA report (authored)", "Narrative insight + metric + recommendation grounded in mart queries"],
            ["Qdrant OTA_insights", "Optional vector corpus so Ask ADZA can retrieve authored OTAs"],
        ],
    )
    _add_body(doc, "Workflow:", bold=True)
    _add_numbered(
        doc,
        [
            "Frame the OTA question (sourcing, prices, climate, food security, trade, …).",
            "Pick fact or aggregate + filters from the recipes / workbook.",
            "Query BigQuery mart_dev with grain and source_key discipline.",
            "Record metric values with units, place, as-of date, and source.",
            "Write insight and recommendation lanes; cite ACF-friendly fields.",
        ],
    )

    _heading(doc, "3. Star-schema overview")
    _add_code(
        doc,
        """dim_geography / dim_product|item / dim_element|unit / dim_source|date|season
        ▲
        │ *_key joins
fct_* (measures + ACF)  ──►  agg_* (country/month rollups)
        │
      bridge_* (AEZ, research)""",
    )
    _add_body(
        doc,
        "Inventory (63 live models): 22 dimensions · 27 facts · 11 aggregates · 3 bridges. "
        "Full list: workbook Entities sheet.",
    )
    _add_bullets(
        doc,
        [
            "Prefer agg_* for national/monthly rollups in OTA metrics.",
            "Drop to fct_* for market, FNID, point, or long-form element detail.",
            "Always keep source_key — never silently blend sources.",
            "Filter grain columns (production_grain, trade_grain, climate_grain, …) before charting.",
        ],
    )

    _heading(doc, "4. Core dimensions")
    _add_table(
        doc,
        ["Dimension", "Use when"],
        [
            ["dim_geography", "Every geo-aware fact — join geography_key; filter geo_level / data_level / country_iso3"],
            ["dim_ref_country", "Africa-scope policy (in_africa_scope); not a substitute for fact→geography joins"],
            ["dim_product", "Crops/commodities on production, prices, food balance, many trade rows"],
            ["dim_item", "FAOSTAT long-form item codes on emissions, employment, land inputs, investment"],
            ["dim_element", "Which measure in long-form facts"],
            ["dim_source", "Lineage + ACF tier — required for trustworthy OTAs"],
            ["dim_date", "Calendar helpers via date_key"],
            ["dim_season", "Seasonal yield — not the same as FAOSTAT year"],
            ["dim_market", "Market-level price facts"],
            ["dim_classification", "FEWS IPC phase labels"],
            ["dim_unit", "Units / currencies"],
        ],
    )
    _add_body(
        doc,
        "Rule: Prefer dim_product for commodity OTAs; use dim_item when the fact only exposes item_key.",
        bold=True,
    )

    _heading(doc, "5. Fact catalog by domain")
    _heading(doc, "Production & inputs", level=2)
    _add_table(
        doc,
        ["Table", "Grain / filter", "OTA use"],
        [
            ["fct_production", "production_grain (physical|index|gross_value)", "National production / area / value"],
            ["fct_yield", "FNID–season yield_raw", "Seasonal yield risk"],
            ["fct_land_inputs", "input_grain", "Fertilizer / land use intensity"],
            ["fct_forestry", "forestry_grain", "Forestry supply"],
            ["fct_machinery", "Country–year", "Mechanisation context"],
        ],
    )
    _heading(doc, "Prices & trade", level=2)
    _add_table(
        doc,
        ["Table", "Grain / filter", "OTA use"],
        [
            ["fct_prices", "Market × month", "Market detail"],
            ["fct_trade", "trade_grain", "FAOSTAT trade vs FEWS borders"],
        ],
    )
    _heading(doc, "Food security & balance", level=2)
    _add_table(
        doc,
        ["Table", "Grain / filter", "OTA use"],
        [
            ["fct_food_security", "measure_type; FNID–month", "Early warning"],
            ["fct_food_balance", "Long-form; map by element_code", "Supply utilisation"],
            ["fct_humanitarian", "Country–year", "Aid volumes"],
            ["fct_food_hazards", "Study-level", "Food safety"],
        ],
    )
    _heading(doc, "Climate, vegetation, socio-economic, spatial", level=2)
    _add_bullets(
        doc,
        [
            "fct_climate — filter climate_grain (point_obs | country_model).",
            "fct_emissions / fct_vegetation / fct_air_quality — GHG, NDVI/ILRI sites, local PM.",
            "fct_soil_health — prefer agg_soil_admin for summaries.",
            "fct_economics / fct_employment / fct_hdi / fct_gender_inclusion — macro context; keep national gender separate from household.",
            "fct_household / fct_animal_health / fct_insurance — community livelihood OTAs.",
            "fct_investment — ASTI / ag R&D funding.",
            "fct_protected_areas / fct_germplasm / fct_biodiversity — conservation / genetic resources (verify geo coverage).",
        ],
    )

    _heading(doc, "6. Aggregates — when to prefer them")
    _add_table(
        doc,
        ["Aggregate", "Prefer for"],
        [
            ["agg_production_country_year", "Cross-country production OTAs (physical)"],
            ["agg_production_country_season", "Seasonal country rollups from yield"],
            ["agg_prices_country_month", "National price trends"],
            ["agg_food_security_country_month", "National FEWS population-in-phase"],
            ["agg_food_balance_country_year", "National food/feed/losses sums"],
            ["agg_economics / employment / emissions / forestry _country_year", "Domain country-year averages"],
            ["agg_hdi_latest", "Latest HDI snapshot (not a time series)"],
            ["agg_soil_admin", "Admin soil summaries vs millions of points"],
        ],
    )
    _add_body(doc, "Aggregates still carry source_key — keep it in GROUP BY / filters.", bold=True)

    _heading(doc, "7. ACF and citation hygiene")
    _add_body(doc, "Every trustworthy OTA metric should carry:")
    _add_bullets(
        doc,
        [
            "What: metric, value, unit",
            "Where: country_iso3, place_scope / geography attributes",
            "When: as_of_date (+ note as_of_date_basis if loaded_at)",
            "Who/source: source_key / source_id, tier",
        ],
    )
    _add_body(
        doc,
        "See workbook sheet ACF_Contract and CATALOG_TO_MART_MAP.md. Do not invent units or fill null WFP common_unit_price.",
    )

    _heading(doc, "8. Hard grain rules")
    _add_numbered(
        doc,
        [
            "fct_production ≠ fct_yield — never average across them; state which grain the OTA uses.",
            "production_grain: use physical for area/qty aggregates (agg_production_country_year).",
            "Price months must be calendar 1–12 (or null).",
            "Always group by source_key — no silent source blend.",
            "Filter grain discriminators before charts: trade_grain, climate_grain, vegetation_grain, forestry_grain, input_grain, measure_type.",
            "Food balance food/feed/losses map by element_code, not English labels alone.",
            "National fct_gender_inclusion stays separate from household gender scores.",
            "ILRI dairy cow-day metrics are out of production/yield facts by design.",
            "Partition pruning: filter as_of_date with ranges — avoid EXTRACT(year FROM as_of_date) when possible.",
        ],
    )

    _heading(doc, "9. OTA report recipes")
    recipes = [
        (
            "Cross-country production / sourcing risk",
            "agg_production_country_year + dim_geography + dim_product + dim_source. "
            "Filter products and years; keep source_key. Metric + insight + recommendation.",
        ),
        (
            "Price volatility & market brief",
            "agg_prices_country_month (national) or fct_prices + dim_market. "
            "Disclose WFP/FEWS geo coverage gaps; months 1–12 only.",
        ),
        (
            "Climate stress vs yield",
            "fct_climate (filter climate_grain) + fct_yield or agg_production_country_season. "
            "Align on country_iso3 + year/season — do not blend FAOSTAT production into seasonal yield.",
        ),
        (
            "Food security early warning",
            "fct_food_security or agg_food_security_country_month + dim_classification. "
            "pct_phase3/4/5 typically on population product only.",
        ),
        (
            "Trade / border flows",
            "fct_trade filtered by trade_grain. FEWS: source_country, destination_country, border_point.",
        ),
    ]
    for title_t, body in recipes:
        _heading(doc, title_t, level=2)
        _add_body(doc, body)

    _heading(doc, "10. SQL patterns")
    _add_body(doc, "National production (physical):", bold=True)
    _add_code(
        doc,
        """SELECT g.country_iso3, p.product_name, a.year,
       a.area_harvested, a.production_qty, a.yield_recomputed, a.source_key
FROM `opentrace-prod-5ga4.mart_dev.agg_production_country_year` a
JOIN `opentrace-prod-5ga4.mart_dev.dim_geography` g USING (geography_key)
JOIN `opentrace-prod-5ga4.mart_dev.dim_product` p USING (product_key)
WHERE g.country_iso3 IN ('NGA', 'GHA', 'CIV')
  AND LOWER(p.product_name) LIKE '%maize%'
  AND a.year BETWEEN 2019 AND 2024;""",
    )
    _add_body(doc, "Country-month prices:", bold=True)
    _add_code(
        doc,
        """SELECT g.country_iso3, pr.product_name, a.year, a.month,
       a.price_avg, a.common_unit_price_avg, a.source_key
FROM `opentrace-prod-5ga4.mart_dev.agg_prices_country_month` a
JOIN `opentrace-prod-5ga4.mart_dev.dim_geography` g USING (geography_key)
JOIN `opentrace-prod-5ga4.mart_dev.dim_product` pr USING (product_key)
WHERE g.country_iso3 IN ('ETH', 'KEN') AND a.year = 2024;""",
    )
    _add_body(
        doc,
        "Replace project id if your environment differs; dataset name remains mart_dev.",
        italic=True,
    )

    _heading(doc, "11. Known caveats")
    _add_bullets(
        doc,
        [
            "HDI geo historically problematic — verify geography_key / country_iso3 before citing.",
            "Prices: WFP/FEWS geo_null can be high; WFP common_unit_price null is expected.",
            "Soil iSDA: large volume; weak as_of historically; prefer agg_soil_admin.",
            "Point facts: check geo coverage; nearest-city resolution applies.",
            "Cluster-only facts may not prune on as_of_date partitions.",
            "FEWS: null phase % on classification rows is expected.",
        ],
    )
    _add_body(doc, "Details: MART_QA_NOTES.md")

    _heading(doc, "12. Appendix — regeneration")
    _add_bullets(
        doc,
        [
            "Workbook: data-eng/docs/mart_dev_entity_dictionary.xlsx",
            "Seed: data-eng/docs/mart_entity_dictionary_seed.yaml",
            "Markdown twin: data-eng/docs/MART_DEV_OTA_ANALYST_GUIDE.md",
        ],
    )
    _add_code(
        doc,
        """cd data-eng
python scripts/build_mart_entity_dictionary.py
python scripts/build_mart_ota_analyst_guide_docx.py""",
    )
    _add_body(doc, "Contact: contact@opentrace.africa", bold=True)
    return doc


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--out", type=Path, default=DEFAULT_OUT)
    args = parser.parse_args()
    args.out.parent.mkdir(parents=True, exist_ok=True)
    build_document().save(str(args.out))
    print(f"Wrote {args.out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
