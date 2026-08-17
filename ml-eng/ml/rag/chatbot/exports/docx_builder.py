"""Build Word (.docx) reports with citations and ACF summary."""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Inches, Pt

from ml.rag.chatbot.exports.markdown_flow import iter_inline_runs, iter_markdown_blocks


def _add_runs(paragraph: Any, text: str) -> None:
    for kind, chunk in iter_inline_runs(text):
        run = paragraph.add_run(chunk)
        run.font.size = Pt(11)
        if kind == "bold":
            run.bold = True
        elif kind == "italic":
            run.italic = True


def _add_markdown_body(doc: Any, body: str) -> None:
    for kind, content in iter_markdown_blocks(body):
        if kind == "heading":
            doc.add_heading(content, level=2)
            continue
        if kind == "bullet":
            paragraph = doc.add_paragraph(style="List Bullet")
            _add_runs(paragraph, content)
            continue
        if kind == "ordered":
            paragraph = doc.add_paragraph(style="List Number")
            _add_runs(paragraph, content)
            continue
        paragraph = doc.add_paragraph()
        _add_runs(paragraph, content)


def build_docx(
    *,
    title: str,
    sections: list[dict[str, str]],
    table_rows: list[dict[str, Any]] | None = None,
    chart_png: bytes | None = None,
    citations: list[dict[str, Any]] | None = None,
    acf_summary: str | None = None,
    filename: str = "report.docx",
) -> tuple[bytes, str]:
    doc = Document()
    doc.add_heading(title, level=0)

    for sec in sections:
        heading = (sec.get("heading") or "").strip()
        body = (sec.get("body") or "").strip()
        if heading:
            doc.add_heading(heading, level=1)
        if body:
            _add_markdown_body(doc, body)

    if chart_png:
        doc.add_heading("Chart", level=1)
        pic_stream = io.BytesIO(chart_png)
        doc.add_picture(pic_stream, width=Inches(6))

    if table_rows:
        doc.add_heading("Figures", level=1)
        cols = list({k for row in table_rows for k in row})
        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(cols):
            hdr[i].text = str(col)
        for row in table_rows[:50]:
            cells = table.add_row().cells
            for i, col in enumerate(cols):
                cells[i].text = str(row.get(col, ""))

    if citations:
        doc.add_heading("Sources", level=1)
        for c in citations:
            cid = c.get("id", "")
            text = c.get("text", "")
            doc.add_paragraph(f"[{cid}] {text}", style="List Bullet")

    if acf_summary:
        doc.add_heading("Confidence summary", level=1)
        doc.add_paragraph(acf_summary)

    buf = io.BytesIO()
    doc.save(buf)
    name = filename if filename.endswith(".docx") else f"{filename}.docx"
    return buf.getvalue(), name


__all__ = ["build_docx"]
