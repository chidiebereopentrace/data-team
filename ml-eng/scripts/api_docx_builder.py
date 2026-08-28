"""Reusable python-docx helpers for OpenTrace API documentation."""
from __future__ import annotations

from datetime import date
from typing import Any, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor


def new_document() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    return doc


def add_title_page(
    doc: Document,
    *,
    title: str,
    subtitle: str,
    version: str,
    generated: date | None = None,
) -> None:
    when = generated or date.today()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(subtitle)
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Version {version}  |  Generated {when.isoformat()}").font.size = Pt(10)

    doc.add_page_break()


def add_toc_placeholder(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Update field in Word: References → Table of Contents")
    run.italic = True
    run.font.size = Pt(10)
    _insert_toc_field(doc)
    doc.add_page_break()


def _insert_toc_field(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, *, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_bullets(doc: Document, items: Sequence[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)


def add_code_block(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)


def add_flowchart_diagram(doc: Document, title: str, lines: Sequence[str]) -> None:
    """Render a flowchart as a monospace block with optional Unicode box-drawing."""
    if title:
        doc.add_heading(title, level=3)
    add_code_block(doc, "\n".join(lines))


def add_entity_table(
    doc: Document,
    entity_name: str,
    fields: Sequence[Sequence[str]],
) -> None:
    """One ERD entity: field name, type, key (PK/FK/empty)."""
    doc.add_heading(entity_name, level=4)
    add_field_table(doc, ("Field", "Type", "Key"), fields)


def add_relationship_table(
    doc: Document,
    title: str,
    rows: Sequence[Sequence[str]],
) -> None:
    """ERD relationships: from_entity, cardinality, to_entity, label."""
    doc.add_heading(title, level=3)
    add_field_table(doc, ("From", "Cardinality", "To", "Label"), rows)


def add_node_inventory_table(
    doc: Document,
    rows: Sequence[Sequence[str]],
) -> None:
    """LangGraph node inventory: node, type, description."""
    add_field_table(doc, ("Node", "Type", "Description"), rows)


def add_erd_section(
    doc: Document,
    *,
    title: str,
    entities: Sequence[tuple[str, Sequence[Sequence[str]]]],
    relationships: Sequence[Sequence[str]],
) -> None:
    """Full ERD subsection: entity attribute tables + relationship matrix."""
    doc.add_heading(title, level=2)
    for entity_name, fields in entities:
        add_entity_table(doc, entity_name, fields)
    add_relationship_table(doc, "Relationships", relationships)


def add_json_example(doc: Document, title: str, json_text: str) -> None:
    doc.add_heading(title, level=3)
    add_code_block(doc, json_text)


def add_field_table(
    doc: Document,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()


def add_endpoint_section(
    doc: Document,
    *,
    method: str,
    path: str,
    description: str,
    request_fields: Sequence[Sequence[str]] | None = None,
    response_fields: Sequence[Sequence[str]] | None = None,
    curl_example: str | None = None,
    request_json: str | None = None,
    response_json: str | None = None,
    errors: Sequence[Sequence[str]] | None = None,
) -> None:
    doc.add_heading(f"{method} {path}", level=2)
    add_paragraph(doc, description)
    if request_fields:
        doc.add_heading("Request fields", level=3)
        add_field_table(doc, ("Field", "Type", "Required", "Description"), request_fields)
    if response_fields:
        doc.add_heading("Response fields", level=3)
        add_field_table(doc, ("Field", "Type", "Description"), response_fields)
    if errors:
        doc.add_heading("Errors", level=3)
        add_field_table(doc, ("Status", "When"), errors)
    if curl_example:
        add_json_example(doc, "cURL example", curl_example)
    if request_json:
        add_json_example(doc, "Request example", request_json)
    if response_json:
        add_json_example(doc, "Response example", response_json)


def add_comparison_table(
    doc: Document,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> None:
    doc.add_heading(title, level=2)
    add_field_table(doc, headers, rows)


def add_architecture_diagram(doc: Document) -> None:
    doc.add_heading("Architecture", level=2)
    diagram = """
Client
  |
  v
RAG API (POST /query)  OR  Chatbot API v1 (POST /v1/chat)
  |
  v
execute_chat_turn()
  |
  v
run_rag()  -- LangGraph pipeline
  |-- Query decomposition
  |-- BigQuery retrieval (staging_dev)
  |-- Qdrant vector retrieval (news, research, policy, OTA, ...)
  |-- Merge + rerank
  |-- LLM generation
  |
  v
Response: answer / assistant_message, citations[], acf, usage, session_id
"""
    add_code_block(doc, diagram.strip())


def save_document(doc: Document, path: str) -> None:
    doc.save(path)
