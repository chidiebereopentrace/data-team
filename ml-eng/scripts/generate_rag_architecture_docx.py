#!/usr/bin/env python3
"""Generate OpenTrace RAG pipeline architecture documentation (Word .docx)."""
from __future__ import annotations

from datetime import date
from typing import Any

from api_docx_builder import (
    add_bullets,
    add_erd_section,
    add_field_table,
    add_flowchart_diagram,
    add_heading,
    add_node_inventory_table,
    add_paragraph,
    add_title_page,
    add_toc_placeholder,
    new_document,
    save_document,
)
from docx.document import Document
from rag_architecture_content import (
    ACF_SUBFLOW_CHART,
    BQ_ENRICH_BULLETS,
    BQ_RETRIEVAL_BULLETS,
    BQ_SUBFLOW_CHART,
    CORPUS_TABLE,
    DOCS_DIR,
    ENTRY_POINTS_FLOWCHART,
    ENV_VARS,
    FULL_GRAPH_FLOWCHART,
    INFRA_ERD_ENTITIES,
    INFRA_ERD_RELATIONSHIPS,
    INGEST_ARCHITECTURE_BULLETS,
    INGEST_ERD_ENTITIES,
    INGEST_ERD_RELATIONSHIPS,
    LLM_MATRIX,
    NODE_INVENTORY,
    NODE_SPECS,
    OBSERVABILITY_BULLETS,
    OUTPUT_DOCX,
    PURPOSE_BULLETS,
    RAG_STATE_FIELDS,
    ROUTING_TABLE,
    RUNTIME_ERD_ENTITIES,
    RUNTIME_ERD_RELATIONSHIPS,
    VECTOR_RETRIEVAL_BULLETS,
)


def _add_node_section(doc: Document, spec: dict[str, Any]) -> None:
    add_heading(doc, f"Node: {spec['name']}", level=2)
    add_paragraph(doc, spec["purpose"])
    add_field_table(
        doc,
        ("Attribute", "Value"),
        [
            ("Module(s)", spec["module"]),
            ("Reads state", spec["reads"]),
            ("Writes state", spec["writes"]),
            ("LLM", spec["llm"]),
            ("External I/O", spec["external"]),
            ("Failure behavior", spec["failure"]),
        ],
    )
    add_heading(doc, "Sub-steps", level=3)
    add_bullets(doc, spec["substeps"])


def build_architecture_document() -> None:
    doc = new_document()
    add_title_page(
        doc,
        title="OpenTrace RAG Pipeline Architecture",
        subtitle="Full LangGraph pipeline, data flows, and node reference",
        version="1.0",
        generated=date.today(),
    )
    add_toc_placeholder(doc)

    add_heading(doc, "Purpose and design principles", level=1)
    add_bullets(doc, PURPOSE_BULLETS)

    add_heading(doc, "Entry points", level=1)
    add_flowchart_diagram(doc, "Entry points flowchart", ENTRY_POINTS_FLOWCHART)

    add_heading(doc, "Full LangGraph (16 nodes)", level=1)
    add_paragraph(doc, "Authoritative routing from chatbot/graph.py build_graph().")
    add_flowchart_diagram(doc, "Runtime graph", FULL_GRAPH_FLOWCHART)
    add_heading(doc, "Node inventory", level=2)
    add_node_inventory_table(doc, NODE_INVENTORY)

    add_heading(doc, "Conditional routing table", level=2)
    add_field_table(doc, ("Stage", "Condition", "Target node", "Notes"), ROUTING_TABLE)

    add_heading(doc, "Entity-relationship diagrams", level=1)

    add_erd_section(
        doc,
        title="Runtime domain (§3.1)",
        entities=RUNTIME_ERD_ENTITIES,
        relationships=RUNTIME_ERD_RELATIONSHIPS,
    )
    add_erd_section(
        doc,
        title="Ingest and vector store (§3.2)",
        entities=INGEST_ERD_ENTITIES,
        relationships=INGEST_ERD_RELATIONSHIPS,
    )
    add_erd_section(
        doc,
        title="Infrastructure (§3.3)",
        entities=INFRA_ERD_ENTITIES,
        relationships=INFRA_ERD_RELATIONSHIPS,
    )

    add_heading(doc, "BQ retrieve sub-pipeline", level=2)
    add_flowchart_diagram(doc, "BQ retrieve flow", BQ_SUBFLOW_CHART)

    add_heading(doc, "Generation and ACF sub-pipeline", level=2)
    add_flowchart_diagram(doc, "Generate + ACF flow", ACF_SUBFLOW_CHART)

    add_heading(doc, "BQ enrich and ACF metadata", level=2)
    add_bullets(doc, BQ_ENRICH_BULLETS)

    add_heading(doc, "Node reference", level=1)
    add_paragraph(doc, "Detailed behavior for each LangGraph node.")
    for spec in NODE_SPECS:
        _add_node_section(doc, spec)

    add_heading(doc, "RAGGraphState field catalog", level=1)
    add_field_table(doc, ("Field", "Type", "Description"), RAG_STATE_FIELDS)

    add_heading(doc, "Retrieval subsystems", level=1)
    add_heading(doc, "Vector retrieval", level=2)
    add_bullets(doc, VECTOR_RETRIEVAL_BULLETS)
    add_heading(doc, "Corpus router (six collections)", level=2)
    add_field_table(doc, ("Key", "Qdrant collection", "Role", "Key payload indexes"), CORPUS_TABLE)
    add_heading(doc, "BigQuery retrieval", level=2)
    add_bullets(doc, BQ_RETRIEVAL_BULLETS)

    add_heading(doc, "Ingest architecture (offline)", level=1)
    add_bullets(doc, INGEST_ARCHITECTURE_BULLETS)

    add_heading(doc, "LLM usage matrix", level=1)
    add_field_table(doc, ("Step", "Module", "Backend", "Notes"), LLM_MATRIX)

    add_heading(doc, "Observability and debugging", level=1)
    add_bullets(doc, OBSERVABILITY_BULLETS)

    add_heading(doc, "Environment variables (subset)", level=1)
    add_field_table(doc, ("Variable", "Purpose"), ENV_VARS)

    add_heading(doc, "Related documentation", level=1)
    add_bullets(
        doc,
        [
            "OpenTrace-RAG-Pipeline-Architecture.pdf — visual diagrams (Mermaid)",
            "OpenTrace-RAG-API-Documentation.docx — HTTP /query reference",
            "OpenTrace-Chatbot-API-v1-Documentation.docx — public v1 chat + artifacts",
            "ml/rag/docs/API.md — markdown API reference",
            "ml/rag/ARCHITECTURE.md — this package architecture (markdown)",
            "ml/rag/docs/SCRIPTS.md — CLI and ingestion commands",
            "Regenerate DOCX: python scripts/generate_rag_architecture_docx.py from ml-eng/",
            "Regenerate PDF: python scripts/generate_rag_architecture_pdf.py from ml-eng/",
        ],
    )

    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    save_document(doc, str(OUTPUT_DOCX))
    print(f"Wrote {OUTPUT_DOCX}")


def main() -> None:
    build_architecture_document()


if __name__ == "__main__":
    main()
