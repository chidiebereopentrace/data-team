"""Build CEO-shareable Word document for Olam AGRI Enterprise API pre-sales brief."""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import _Cell

REPO_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_OUT = REPO_ROOT / "docs" / "partners" / "ASKADZA_ENTERPRISE_API_OLAM_AGRI.docx"


def _shade_header_row(table) -> None:
    for cell in table.rows[0].cells:
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = tc_pr.makeelement(qn("w:shd"), {})
            tc_pr.append(shd)
        shd.set(qn("w:fill"), "E8E8E8")


def _set_cell_text(cell: _Cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _add_body(doc: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic


def _add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.size = Pt(11)


def _add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        run.font.size = Pt(11)


def _add_code(doc: Document, text: str) -> None:
    for line in text.strip().splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(0)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[i], header, bold=True)
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            _set_cell_text(table.rows[r_idx + 1].cells[c_idx], value)
    _shade_header_row(table)
    doc.add_paragraph()


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def build_document() -> Document:
    doc = Document()

    title = doc.add_heading("Ask ADZA Enterprise API", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.size = Pt(22)

    subtitle = doc.add_paragraph("Product Integration Brief for Olam AGRI")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(14)
        run.italic = True

    doc.add_paragraph()
    meta_lines = [
        "Prepared by: OpenTrace",
        "Date: August 2026",
        "Classification: Confidential — for Olam AGRI partnership discussions",
        "Contact: contact@opentrace.africa | opentrace.africa | askadza.africa",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(10)

    doc.add_page_break()

    # Executive summary
    _add_heading(doc, "Executive summary")
    _add_body(
        doc,
        "OpenTrace is Africa's agricultural intelligence layer — harmonising fragmented climate, "
        "production, market, and food-security data across 54 countries into decision-ready intelligence. "
        "Ask ADZA is the natural-language interface; the Enterprise API lets institutions like Olam AGRI "
        "embed that intelligence directly into procurement systems, risk dashboards, and internal tools — "
        "without building data pipelines or hiring analysts in the middle.",
    )
    _add_body(doc, "For Olam AGRI, the Enterprise API delivers:", bold=True)
    _add_bullets(
        doc,
        [
            "Cross-country supply and production intelligence for sourcing and portfolio decisions",
            "Market volatility and price trend analysis across African origins",
            "Climate and yield risk signals at national and sub-national levels",
            "Traceable, confidence-scored answers (ACF) with citations to underlying datasets",
            "Exportable artifacts (CSV, charts, DOCX, PDF) for reports and workflows",
        ],
    )
    _add_body(
        doc,
        '"Individuals use Ask ADZA; institutions connect via API." — OpenTrace commercial model',
        italic=True,
    )

    # Problem
    _add_heading(doc, "The problem Olam AGRI faces")
    _add_body(
        doc,
        "Africa is central to global agri-supply chains, yet decision-makers often lack a unified view of:",
    )
    _add_bullets(
        doc,
        [
            "Production trends and yield stability across origins",
            "Climate variability and its impact on supply reliability",
            "Regional price dynamics and cross-border trade flows",
            "Emerging risks in specific districts and value chains",
        ],
    )
    _add_body(
        doc,
        "Data exists across governments, research bodies, and market systems — but it is fragmented, "
        "inconsistent, and not queryable at decision speed. OpenTrace closes that gap.",
    )

    # What OpenTrace provides
    _add_heading(doc, "What OpenTrace provides")
    _add_body(doc, "OpenTrace is an infrastructure company, not a consultancy. Three pillars:")
    _add_table(
        doc,
        ["Pillar", "Role"],
        [
            [
                "OFIA (OpenTrace Federated Intelligence Architecture)",
                "Harmonises and version-controls datasets across global, national, and sub-national levels",
            ],
            [
                "ACF (ADZA Confidence Framework)",
                "Attaches honest confidence signals to every answer based on evidence triangulation",
            ],
            [
                "Ask ADZA",
                "Natural-language interface — web, mobile, WhatsApp, and Enterprise API",
            ],
        ],
    )
    _add_body(
        doc,
        "Coverage today: 2B+ data points | 54 African countries | 12+ indicator domains "
        "(production, prices, trade, climate, food security, soil, employment, and more)",
        bold=True,
    )
    _add_body(doc, "Data domains relevant to Olam AGRI:", bold=True)
    _add_bullets(
        doc,
        [
            "Crop production and yield trends (FAOSTAT, sub-national yield where available)",
            "Market prices and food balance sheets",
            "Cross-border trade flows",
            "Climate and vegetation indices (ERA5, NDVI, NASA POWER)",
            "Food security early warning (FEWS NET)",
            "Employment and macro indicators for market context",
        ],
    )

    # Enterprise API
    _add_heading(doc, "Enterprise API — how it works")
    _add_body(
        doc,
        "The Enterprise API is a RESTful, versioned HTTP service that accepts natural-language "
        "questions and returns structured, cited intelligence.",
    )
    _add_body(doc, "Integration pattern:", bold=True)
    _add_code(
        doc,
        """Olam application  →  POST /v1/chat/agribusinesses  →  OpenTrace intelligence engine
                     (API key in header)              →  Answer + citations + ACF + optional exports""",
    )
    _add_body(doc, "Authentication: Send your API key on every request:", bold=True)
    _add_code(doc, "X-API-Key: <your-tenant-api-key>")
    _add_body(doc, "Or:")
    _add_code(doc, "Authorization: Bearer <your-tenant-api-key>")
    _add_body(doc, "Key capabilities for agribusiness partners:", bold=True)
    _add_table(
        doc,
        ["Capability", "Description"],
        [
            ["Natural-language Q&A", "Ask questions in plain English; no SQL or dashboards required"],
            [
                "Multi-turn conversations",
                "Server-side sessions or client-owned history for copilot-style UX",
            ],
            [
                "Cross-country comparison",
                "Compare production, prices, or climate across African origins",
            ],
            [
                "Structured citations",
                "Every claim links to source datasets (news, research, structured data)",
            ],
            [
                "ACF confidence scoring",
                "Transparent strong / moderate / limited bands with explanation",
            ],
            ["Export artifacts", "CSV, chart (PNG), DOCX, and PDF downloads via signed URLs"],
            [
                "Usage metering",
                "Token usage returned per request; monthly quotas enforced per tenant",
            ],
        ],
    )
    _add_body(doc, "Example questions Olam teams could ask via API:", bold=True)
    _add_bullets(
        doc,
        [
            '"Compare maize production trends in Nigeria, Ghana, and Côte d\'Ivoire over the last five years."',
            '"Which regions in East Africa show the highest climate stress on coffee yields?"',
            '"What are retail maize price trends in Ethiopia and Kenya this season?"',
            '"Export a CSV of rice trade flows for West African countries in 2024."',
        ],
    )

    # Integration options
    _add_heading(doc, "Integration options for Olam AGRI")
    _add_table(
        doc,
        ["Option", "Description", "Best for"],
        [
            [
                "Embedded copilot",
                "API-backed chat widget in procurement or risk portals",
                "Analysts and category managers",
            ],
            [
                "Backend intelligence service",
                "Server-to-server calls from Olam's risk/scenario models",
                "Automated monitoring and alerts",
            ],
            [
                "Report generation",
                "Scheduled API queries → export artifacts → internal distribution",
                "Weekly/monthly origin briefings",
            ],
            [
                "WhatsApp / field channel",
                "Ask ADZA answers surfaced to field teams (future channel)",
                "Origin managers and agronomists",
            ],
        ],
    )
    _add_body(
        doc,
        "OpenTrace does not replace Olam's internal systems. Partners plug into shared intelligence "
        "infrastructure — data stays under original licences; OpenTrace monetizes intelligence, not raw data.",
    )

    # Data trust
    _add_heading(doc, "Data trust and sovereignty")
    _add_body(doc, "OpenTrace is built for institutional trust:")
    _add_bullets(
        doc,
        [
            "Partners retain ownership of any data they contribute; attribution is preserved",
            "Source data stays under its original licence — OpenTrace does not relicense third-party datasets",
            "No black-box scores — every answer is interrogable via citations and ACF",
            "Derived intelligence is separate from source data",
            "ACF triangulation across global (25%), national/regional (40%), and ground/community (35%) evidence tiers",
        ],
    )

    # Commercial model
    _add_heading(doc, "Commercial model (indicative)")
    _add_body(doc, "OpenTrace monetizes intelligence access, not data enclosure:")
    _add_table(
        doc,
        ["Component", "Model"],
        [
            ["Platform fee", "Annual enterprise licence for API access"],
            ["Usage", "Per-query or per-token consumption above included quota"],
            [
                "Exports",
                "Included in Agribusinesses/Integrated tier; volume tiers for high export use",
            ],
            [
                "Custom scope",
                "Optional add-ons: additional countries, proprietary data integration, co-branded exports",
            ],
            [
                "Pilot",
                "Reduced-fee or fee-waived sandbox period (typically 4–8 weeks) to validate use cases",
            ],
        ],
    )
    _add_body(
        doc,
        "Specific pricing is scoped during partnership discovery based on query volume, countries, "
        "and integration depth.",
    )

    # Technical requirements
    _add_heading(doc, "Technical requirements (Olam side)")
    _add_body(doc, "Minimal integration footprint:")
    _add_bullets(
        doc,
        [
            "HTTPS client capable of POST requests with JSON body",
            "API key storage in Olam's secrets management (server-side only; never in client apps)",
            "Session handling — reuse session_id for multi-turn, or pass chat_history for stateless calls",
            "Citation rendering — map inline [N] footnotes to citation cards in your UI (optional but recommended)",
            "Export handling — download artifacts from signed GCS URLs returned in artifacts[]",
        ],
    )
    _add_body(
        doc,
        "Full technical reference: Enterprise Integration Guide and OpenAPI spec (available on request).",
    )

    # Partnership journey
    _add_heading(doc, "Proposed partnership journey")
    _add_table(
        doc,
        ["Stage", "Activities", "Duration"],
        [
            [
                "1. Discovery",
                "Map Olam use cases (sourcing, risk, origins); agree success metrics",
                "1–2 weeks",
            ],
            ["2. Sandbox", "API key, test environment, sample integration", "1 week"],
            [
                "3. Pilot",
                "Live integration in 2–3 workflows; joint evaluation of accuracy and citation quality",
                "4–8 weeks",
            ],
            ["4. Production", "Production API, SLA, billing, partner success support", "Ongoing"],
            [
                "5. Expand",
                "Additional use cases, optional proprietary data federation into OFIA",
                "As needed",
            ],
        ],
    )
    _add_body(doc, "Pilot success metrics (examples):", bold=True)
    _add_bullets(
        doc,
        [
            "≥90% of pilot queries return actionable answers with citations",
            "ACF confidence distribution aligns with Olam analyst validation on sample set",
            "Time-to-insight reduced vs. manual data gathering baseline",
            "Export artifacts usable in existing Olam reporting workflows",
        ],
    )

    # Why partner
    _add_heading(doc, "Why partner with OpenTrace")
    _add_table(
        doc,
        ["For Olam AGRI", "OpenTrace delivers"],
        [
            [
                "Sourcing and supply risk",
                "Cross-country production, climate, and market intelligence on demand",
            ],
            [
                "Speed to insight",
                "Natural language instead of manual reconciliation across datasets",
            ],
            ["Institutional trust", "ACF confidence scoring and full citation traceability"],
            [
                "African coverage",
                "54 countries, continental scale, sub-national where data allows",
            ],
            [
                "Low integration burden",
                "REST API — no data warehouse to build or maintain",
            ],
            [
                "Long-term infrastructure",
                "OFIA persists and improves with use; not a one-off consultancy project",
            ],
        ],
    )
    _add_body(
        doc,
        "OpenTrace already collaborates with agricultural research and development institutions "
        "(IITA, CGIAR ecosystem, AGRA, governments, and foundations). Olam AGRI would join a growing "
        "network of institutions leveraging shared agricultural intelligence infrastructure.",
    )

    # Next steps
    _add_heading(doc, "Next steps")
    _add_numbered(
        doc,
        [
            "Introductory call — align on Olam priorities and candidate use cases",
            "NDA + sandbox provisioning — API access for technical evaluation",
            "Discovery workshop — Olam category/risk teams + OpenTrace product/engineering",
            "Pilot agreement — scope, timeline, success criteria, and commercial terms",
            "Integration kickoff — Olam engineering receives API credentials, OpenAPI spec, and integration guide",
        ],
    )
    _add_body(doc, "Contact: contact@opentrace.africa", bold=True)

    # Appendix A
    _add_heading(doc, "Appendix A — API surface summary (Agribusinesses tier)", level=2)
    _add_table(
        doc,
        ["Method", "Endpoint", "Purpose"],
        [
            ["GET", "/v1/health", "Service health"],
            ["GET", "/v1/meta", "Plan types, categories, rate limits"],
            ["GET", "/v1/usage", "Current tenant usage (authenticated)"],
            ["POST", "/v1/sessions", "Create conversation session"],
            ["GET", "/v1/sessions/{id}", "Session status"],
            ["POST", "/v1/chat/agribusinesses", "Primary endpoint — Q&A with exports"],
            ["POST", "/v1/feedback", "Optional quality feedback (when tracing enabled)"],
        ],
    )
    _add_body(doc, "Request body (minimal):", bold=True)
    _add_code(
        doc,
        """{
  "message": "Compare cocoa production in Ghana and Côte d'Ivoire over the last 3 years",
  "session_id": "optional-for-multi-turn"
}""",
    )
    _add_body(
        doc,
        "Response includes: assistant_message, citations[], acf, usage, artifacts[] "
        "(when export requested), session_id, request_id",
    )

    # Appendix B
    _add_heading(doc, "Appendix B — OpenTrace readiness timeline", level=2)
    _add_body(doc, "OpenTrace is completing the following for production B2B launch:")
    _add_bullets(
        doc,
        [
            "API gateway and client authentication — available in sandbox",
            "Per-tenant metering and billing pipeline — available in sandbox",
            "Partner sandbox environment — available on request",
            "Published OpenAPI specification and integration guide — available",
            "Enterprise MSA, DPA, and SLA templates — available under NDA",
        ],
    )
    _add_body(
        doc,
        "Pilot can begin on a controlled sandbox while production gateway hardening completes in parallel.",
    )

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Build Olam AGRI Enterprise API CEO brief DOCX")
    parser.add_argument(
        "--out",
        type=Path,
        default=DEFAULT_OUT,
        help=f"Output path (default: {DEFAULT_OUT})",
    )
    args = parser.parse_args()

    out_path: Path = args.out
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = build_document()
    doc.save(str(out_path))
    print(f"Wrote {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
